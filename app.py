"""
app.py - AuditX 就業規則チェックツール v0.5
"""

import json
import os
import tempfile
from datetime import datetime
from pathlib import Path

import io
import re

import streamlit as st
from audit_engine import PHASE_LABELS, extract_text, generate_report, run_audit, extract_seishain_summary, SEISHAIN_SUMMARY_ITEMS

# ============================================================
# 定数
# ============================================================

COURSES = [
    {"id": "CA_seishain",  "group": "キャリアアップ助成金",   "name": "正社員化コース"},
    {"id": "CA_shoyo",     "group": "キャリアアップ助成金",   "name": "賞与・退職金制度導入コース"},
    {"id": "KO65_keizoku", "group": "65歳超雇用推進助成金",   "name": "65歳超継続雇用促進コース"},
    {"id": "KO65_tenkan",  "group": "65歳超雇用推進助成金",   "name": "高年齢者無期雇用転換コース"},
    {"id": "JK_kanri",     "group": "人材確保等支援助成金",   "name": "雇用管理制度・雇用環境整備助成コース"},
    {"id": "JK_hyoka",     "group": "人材確保等支援助成金",   "name": "人事評価改善等助成コース"},
    {"id": "JH_kyuka",     "group": "人材開発支援助成金",     "name": "教育訓練休暇等付与コース"},
    {"id": "RY_funin",     "group": "両立支援等助成金",       "name": "不妊治療及び女性の健康課題対応両立支援コース"},
    {"id": "RY_juman",     "group": "両立支援等助成金",       "name": "柔軟な働き方選択制度等支援コース"},
    {"id": "RY_shussei",   "group": "両立支援等助成金",       "name": "出生時両立支援コース"},
    {"id": "RY_kaigo",     "group": "両立支援等助成金",       "name": "介護離職防止支援コース"},
    {"id": "RY_ikukyu",    "group": "両立支援等助成金",       "name": "育児休業等支援コース"},
    {"id": "RY_daitai",    "group": "両立支援等助成金",       "name": "育休中等業務代替支援コース"},
]

YEAR_OPTIONS = ["R08", "R07", "R06"]
DATE_OPTIONS = {"R08": ["0408", "0401"], "R07": ["0401"], "R06": ["0401"]}

PHASE_ITEMS = {
    "phase1": ("PHASE 1", "申請準備開始時", "就業規則新規作成後の初回確認"),
    "phase2": ("PHASE 2", "制度導入・改訂時", "導入内容・改訂内容の不備確認"),
    "phase3": ("PHASE 3", "支給申請提出前", "全バージョン整合性の最終確認"),
}

# ============================================================
# ページ設定
# ============================================================

st.set_page_config(
    page_title="就業規則チェックツール | Humax",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+JP:wght@300;400;500;700&display=swap');

html, body, [class*="css"] {
    font-family: 'Noto Sans JP', sans-serif;
    background-color: #dde0e7;
    color: #1e293b;
}
.block-container { padding: 2.5rem 2rem 5rem; max-width: 1100px; }

section[data-testid="stSidebar"] {
    background: #0d1b30;
    border-right: none;
}
section[data-testid="stSidebar"] * { color: #b0bec5 !important; }
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3,
section[data-testid="stSidebar"] strong { color: #e8dfc0 !important; }

.page-eyebrow {
    font-size: 0.68rem; font-weight: 700; letter-spacing: 0.18em;
    color: #c8a84a; text-transform: uppercase; margin-bottom: 0.5rem;
}
.page-title {
    font-size: 2.1rem; font-weight: 700; color: #0f172a;
    margin: 0 0 0.4rem; letter-spacing: -0.5px; line-height: 1.2;
}
.page-subtitle {
    font-size: 0.85rem; color: #64748b; font-weight: 400; margin: 0 0 2rem;
}
.page-divider { border: none; border-top: 2px solid #dde2ea; margin: 0 0 2rem; }

.sec-label {
    font-size: 0.72rem; font-weight: 700; letter-spacing: 0.16em;
    color: #c8a84a; text-transform: uppercase;
    display: flex; align-items: center; gap: 12px;
    margin: 2rem 0 1rem;
}
.sec-label::after {
    content: '';
    flex: 1;
    height: 1.5px;
    background: linear-gradient(to right, #c8a84a55, transparent);
}

.card {
    background: #ffffff; border: 1px solid #dde2ea;
    border-radius: 12px; padding: 1.4rem 1.6rem;
    margin-bottom: 0.8rem;
    box-shadow: 0 2px 6px rgba(0,0,0,0.05);
}

.phase-active {
    background: #fffdf5;
    border: 2.5px solid #c8a84a;
    border-radius: 10px; padding: 0.9rem 1rem; margin-bottom: 6px;
    box-shadow: 0 4px 16px rgba(200,168,74,0.28);
}
.phase-inactive {
    background: #ffffff; border: 1.5px solid #dde2ea;
    border-radius: 10px; padding: 0.9rem 1rem; margin-bottom: 6px;
    box-shadow: 0 1px 3px rgba(0,0,0,0.05);
}
.phase-num-active {
    font-size: 0.68rem; font-weight: 700; color: #a07a28;
    text-transform: uppercase; letter-spacing: 0.14em; margin-bottom: 3px;
}
.phase-num-inactive {
    font-size: 0.68rem; font-weight: 700; color: #b0bec5;
    text-transform: uppercase; letter-spacing: 0.14em; margin-bottom: 3px;
}
.phase-title-active {
    font-size: 0.95rem; font-weight: 800; color: #0d1b30; margin-bottom: 2px;
}
.phase-title-inactive {
    font-size: 0.9rem; font-weight: 500; color: #94a3b8; margin-bottom: 2px;
}
.phase-desc-active  { font-size: 0.75rem; color: #475569; }
.phase-desc-inactive { font-size: 0.75rem; color: #b0bec5; }

.phase-hint {
    background: #faf7ee;
    border: 1px solid #e8dfc0;
    border-left: 4px solid #c8a84a;
    border-radius: 0 8px 8px 0;
    padding: 0.75rem 1.1rem;
    font-size: 0.85rem; color: #5a4a1a; font-weight: 500;
    margin: 0.8rem 0 1.5rem;
}

.group-label {
    font-size: 0.73rem; font-weight: 700; letter-spacing: 0.1em;
    color: #64748b; text-transform: uppercase;
    background: #f0f1f4; border-radius: 4px;
    display: inline-block; padding: 2px 8px;
    margin: 1rem 0 0.5rem;
}

.stButton > button {
    font-family: 'Noto Sans JP', sans-serif !important;
    border-radius: 8px !important;
    font-size: 0.85rem !important;
    font-weight: 500 !important;
    transition: all 0.15s !important;
}

.btn-inactive > div > button {
    background: #f1f3f7 !important;
    color: #64748b !important;
    border: 1.5px solid #dde2ea !important;
    box-shadow: 0 1px 2px rgba(0,0,0,0.05) !important;
}
.btn-inactive > div > button:hover {
    background: #e8eaf0 !important;
    border-color: #c8a84a !important;
    color: #a07a28 !important;
}

.btn-active > div > button,
.btn-active > div > button:focus,
.btn-active > div > button:active,
.btn-active button,
.btn-active [data-testid="stBaseButton-secondary"] {
    background: #0d1b30 !important;
    background-color: #0d1b30 !important;
    color: #c8a84a !important;
    border: 2px solid #c8a84a !important;
    box-shadow: 0 3px 12px rgba(13,27,48,0.35) !important;
    font-weight: 700 !important;
    letter-spacing: 0.05em !important;
}
.btn-active > div > button:hover,
.btn-active button:hover {
    background: #162540 !important;
    background-color: #162540 !important;
    color: #c8a84a !important;
}

.stButton > button[kind="primary"],
[data-testid="stBaseButton-primary"] {
    background: #0d1b30 !important;
    background-color: #0d1b30 !important;
    color: #c8a84a !important;
    border: 2px solid #c8a84a !important;
    border-radius: 10px !important;
    font-size: 1rem !important;
    font-weight: 700 !important;
    letter-spacing: 0.5px !important;
    box-shadow: 0 4px 14px rgba(13,27,48,0.25) !important;
    padding: 0.6rem 1rem !important;
}
.stButton > button[kind="primary"]:hover,
[data-testid="stBaseButton-primary"]:hover {
    background: #162540 !important;
    background-color: #162540 !important;
    color: #c8a84a !important;
    box-shadow: 0 6px 18px rgba(13,27,48,0.35) !important;
    transform: translateY(-1px) !important;
}
.stButton > button:disabled {
    background: #f0f1f4 !important;
    color: #a0aab8 !important;
    box-shadow: none !important;
    border: 1px solid #dde2ea !important;
    transform: none !important;
}

.stDownloadButton > button {
    background: #ffffff !important;
    border: 1.5px solid #c8a84a !important;
    color: #a07a28 !important;
    border-radius: 8px !important;
    font-family: 'Noto Sans JP', sans-serif !important;
    font-weight: 600 !important;
}
.stDownloadButton > button:hover {
    background: #faf7ee !important;
}

.stTextInput > div > div > input {
    background: #ffffff !important;
    border-color: #dde2ea !important;
    border-bottom: 1.5px solid #c8a84a !important;
    color: #1e293b !important;
    border-radius: 8px !important;
    font-family: 'Noto Sans JP', sans-serif !important;
    font-size: 0.9rem !important;
}
.stTextInput > div > div > input:focus {
    border-color: #c8a84a !important;
    box-shadow: 0 0 0 2px rgba(200,168,74,0.15) !important;
}
.stSelectbox > div > div {
    background: #ffffff !important;
    border-color: #dde2ea !important;
    color: #1e293b !important;
    font-family: 'Noto Sans JP', sans-serif !important;
}
.stCheckbox > label {
    color: #1e293b !important; font-size: 0.9rem !important;
    font-family: 'Noto Sans JP', sans-serif !important;
    font-weight: 500 !important;
}

.divider { border: none; border-top: 1.5px solid #dde2ea; margin: 2rem 0; }

.sb-brand {
    padding: 0.8rem 0 1.2rem;
    border-bottom: 1px solid #1e3050; margin-bottom: 1.2rem;
}
.sb-brand-eye {
    font-size: 0.63rem; font-weight: 700; letter-spacing: 0.18em;
    color: #c8a84a; text-transform: uppercase; margin-bottom: 5px;
}
.sb-brand-title {
    font-size: 1rem; font-weight: 700; color: #e8dfc0; margin: 0;
}
.sb-brand-sub { font-size: 0.72rem; color: #607080; margin: 3px 0 0; }
.sb-rule-item {
    background: #0a1525;
    border: 1px solid #c8a84a44;
    border-radius: 7px; padding: 6px 10px;
    font-size: 0.77rem; color: #c8a84a;
    margin-bottom: 5px; display: block;
}
.sb-footer {
    position: fixed; bottom: 1.2rem;
    font-size: 0.68rem; color: #3d5068; text-align: center; width: 200px;
}
</style>
""", unsafe_allow_html=True)


# ============================================================
# ユーティリティ
# ============================================================

@st.cache_data
def load_rule_knowledge(course_id, year, date):
    p = Path(__file__).parent / "rule_knowledge" / f"{year}_{date}_{course_id}.json"
    if p.exists():
        with open(p, "r", encoding="utf-8") as f:
            return json.load(f)
    return None

def get_available_keys():
    rule_dir = Path(__file__).parent / "rule_knowledge"
    keys = set()
    if rule_dir.exists():
        for jf in rule_dir.glob("*.json"):
            parts = jf.stem.split("_", 2)
            if len(parts) == 3:
                keys.add(tuple(parts))
    return keys


# ============================================================
# レポートパーサー
# ============================================================

ALERT_CONFIG = {
    "CRITICAL": {"icon": "🔴", "label": "CRITICAL（即時修正必須）", "color": "#fff0f0", "border": "#e53e3e"},
    "WARNING":  {"icon": "🟡", "label": "WARNING（要修正）",       "color": "#fffbeb", "border": "#d69e2e"},
    "CAUTION":  {"icon": "🟠", "label": "CAUTION（要確認）",       "color": "#fff7ed", "border": "#dd6b20"},
    "HUMAN_CHECK": {"icon": "👤", "label": "HUMAN_CHECK（人間確認）", "color": "#f0f4ff", "border": "#4c6ef5"},
    "OK":       {"icon": "🟢", "label": "OK（問題なし）",           "color": "#f0fff4", "border": "#38a169"},
}

def parse_report(report_text: str) -> list:
    """Markdownレポートをチェック項目リストに変換
    
    AIが出力する形式に対応:
    - ### チェック項目名
    - **判定：CRITICAL 🔴** または 判定：CRITICAL
    - **問題の内容：** テキスト または ❗ 問題の内容　テキスト
    """
    items = []

    def extract_field(body, field_name):
        esc = re.escape(field_name)
        # **フィールド名：** の後から次の **〜：** または --- または末尾まで
        pat = (
            rf'\*\*{esc}[：:]\*\*'
            rf'[ \t]*\n?'
            rf'(.*?)'
            rf'(?=\n\*\*[^\n]+[：:]\*\*|\n---|\Z)'
        )
        m = re.search(pat, body, re.DOTALL)
        if m:
            return m.group(1).strip()
        return ""

    # ブロックを行単位で分割
    lines = report_text.split('\n')
    current_title = None
    current_body_lines = []

    def flush_block(title, body_lines):
        body = '\n'.join(body_lines).strip()
        if not title or not body:
            return None
        level = None
        for key in ALERT_CONFIG:
            if re.search(rf'判定[：:][^\n]{{0,40}}{key}', body):
                level = key
                break
        if level is None:
            return None
        return {
            "title": title,
            "level": level,
            "問題の内容": extract_field(body, "問題の内容"),
            "該当箇所":   extract_field(body, "該当箇所"),
            "修正案":     extract_field(body, "修正案"),
            "審査官の目線": extract_field(body, "審査官の目線"),
            "根拠":       extract_field(body, "根拠"),
            "raw": body,
        }

    for line in lines:
        heading = re.match(r'^#{2,4}\s+(.*)', line)
        if heading:
            if current_title is not None:
                item = flush_block(current_title, current_body_lines)
                if item:
                    items.append(item)
            current_title = heading.group(1).strip()
            current_body_lines = []
        else:
            if current_title is not None:
                current_body_lines.append(line)

    if current_title is not None:
        item = flush_block(current_title, current_body_lines)
        if item:
            items.append(item)

    # パーサーで取れなかった場合はrawをそのまま使う
    if not items:
        # フォールバック：rawテキストをそのまま1項目として返す
        for key in ALERT_CONFIG:
            if key in report_text:
                items.append({
                    "title": "チェック結果",
                    "level": key,
                    "問題の内容": report_text,
                    "該当箇所": "", "修正案": "", "審査官の目線": "", "根拠": "",
                    "raw": report_text,
                })
                break

    return items


def display_report_visual(report_text: str, label: str):
    """チェック結果を視覚的に表示"""
    items = parse_report(report_text)

    if not items:
        st.markdown(report_text)
        return

    # サマリーカウント
    counts = {k: 0 for k in ALERT_CONFIG}
    for it in items:
        if it["level"] in counts:
            counts[it["level"]] += 1

    # サマリーバー
    st.markdown("#### チェックサマリー")
    cols = st.columns(5)
    for col, (key, cfg) in zip(cols, ALERT_CONFIG.items()):
        with col:
            st.markdown(
                f'<div style="background:{cfg["color"]};border:2px solid {cfg["border"]};'
                f'border-radius:10px;padding:0.6rem;text-align:center;">'
                f'<div style="font-size:1.4rem">{cfg["icon"]}</div>'
                f'<div style="font-size:1.3rem;font-weight:800;color:{cfg["border"]}">{counts[key]}</div>'
                f'<div style="font-size:0.65rem;color:#555">{key}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )

    st.markdown("<br>", unsafe_allow_html=True)

    # 優先順位フィルター
    priority_order = ["CRITICAL", "WARNING", "CAUTION", "HUMAN_CHECK", "OK"]
    filter_options = ["すべて表示"] + [
        f'{ALERT_CONFIG[k]["icon"]} {k} ({counts[k]}件)'
        for k in priority_order if counts[k] > 0
    ]
    selected_filter = st.selectbox(
        "表示する判定レベル",
        filter_options,
        key=f"filter_{label[:10]}",
        label_visibility="collapsed",
    )

    # 項目カード表示
    for it in items:
        cfg = ALERT_CONFIG.get(it["level"], {"icon": "ℹ️", "color": "#f8f9fa", "border": "#aaa", "label": it["level"]})
        if selected_filter != "すべて表示" and it["level"] not in selected_filter:
            continue

        with st.expander(f'{cfg["icon"]} {it["title"]}', expanded=(it["level"] in ["CRITICAL", "WARNING"])):
            st.markdown(
                f'<div style="display:inline-block;background:{cfg["color"]};'
                f'border:1.5px solid {cfg["border"]};border-radius:6px;'
                f'padding:3px 12px;font-weight:700;font-size:0.82rem;margin-bottom:0.8rem">'
                f'{cfg["icon"]} {cfg["label"]}</div>',
                unsafe_allow_html=True,
            )
            if it["問題の内容"]:
                st.markdown(f"**❗ 問題の内容**")
                st.markdown(it["問題の内容"])
            if it["該当箇所"]:
                st.markdown(f"**📌 該当箇所**")
                st.code(it["該当箇所"], language=None)
            if it["修正案"]:
                st.markdown(f"**✏️ 修正案**")
                st.info(it["修正案"])
            if it["審査官の目線"]:
                st.markdown(f"**👁️ 審査官の目線**")
                st.markdown(it["審査官の目線"])
            if it["根拠"]:
                st.caption(f"📎 根拠：{it['根拠']}")


# ============================================================
# Excel生成
# ============================================================

def build_xlsx(report_text: str, company: str, label: str, phase: str) -> bytes:
    try:
        import openpyxl
        from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    except ImportError:
        return b""

    items = parse_report(report_text)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "チェック結果"

    # ヘッダー
    headers = ["No.", "チェック項目", "判定", "問題の内容", "該当箇所", "修正案", "根拠"]
    header_fill = PatternFill("solid", fgColor="0D1B30")
    header_font = Font(bold=True, color="C8A84A", size=11)
    col_widths = [5, 30, 12, 50, 40, 50, 30]

    for ci, (h, w) in enumerate(zip(headers, col_widths), 1):
        cell = ws.cell(row=1, column=ci, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.column_dimensions[cell.column_letter].width = w
    ws.row_dimensions[1].height = 25

    # 判定レベル別カラー
    level_colors = {
        "CRITICAL":   "FFE4E4",
        "WARNING":    "FFFBE6",
        "CAUTION":    "FFF3E0",
        "HUMAN_CHECK":"EEF2FF",
        "OK":         "F0FFF4",
    }
    thin = Side(style="thin", color="DDDDDD")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for ri, it in enumerate(items, 2):
        color = level_colors.get(it["level"], "FFFFFF")
        fill = PatternFill("solid", fgColor=color)
        row_data = [
            ri - 1,
            it["title"],
            f'{ALERT_CONFIG.get(it["level"], {}).get("icon","")}{it["level"]}',
            it["問題の内容"],
            it["該当箇所"],
            it["修正案"],
            it["根拠"],
        ]
        for ci, val in enumerate(row_data, 1):
            cell = ws.cell(row=ri, column=ci, value=val)
            cell.fill = fill
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
        ws.row_dimensions[ri].height = 60

    # メタシート
    ws2 = wb.create_sheet("概要")
    ws2["A1"] = "会社名"
    ws2["B1"] = company
    ws2["A2"] = "対象コース"
    ws2["B2"] = label
    ws2["A3"] = "チェックフェーズ"
    ws2["B3"] = PHASE_LABELS.get(phase, phase)
    ws2["A4"] = "出力日時"
    ws2["B4"] = datetime.now().strftime("%Y年%m月%d日 %H:%M")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_xlsx_with_summary(report_text: str, company: str, label: str,
                             phase: str, summary: dict) -> bytes:
    """サマリーシート付きExcel生成（正社員化コース専用）"""
    try:
        import openpyxl
        from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    except ImportError:
        return build_xlsx(report_text, company, label, phase)

    # 通常のチェック結果シートを先に生成
    base_bytes = build_xlsx(report_text, company, label, phase)
    buf_in = io.BytesIO(base_bytes)
    wb = openpyxl.load_workbook(buf_in)

    # サマリーシートを追加
    ws_sum = wb.create_sheet("就業規則サマリー", 0)  # 先頭に挿入

    # ヘッダー
    header_fill = PatternFill("solid", fgColor="0D1B30")
    header_font = Font(bold=True, color="C8A84A", size=11)
    gold_fill   = PatternFill("solid", fgColor="FFF8E7")
    red_fill    = PatternFill("solid", fgColor="FFE4E4")
    thin = Side(style="thin", color="DDDDDD")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    headers = ["項目", "条文番号", "就業規則の記載内容", "⚠️要確認", "実態確認欄（手書き）"]
    col_widths = [22, 14, 55, 10, 30]
    for ci, (h, w) in enumerate(zip(headers, col_widths), 1):
        cell = ws_sum.cell(row=1, column=ci, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        ws_sum.column_dimensions[cell.column_letter].width = w
    ws_sum.row_dimensions[1].height = 25

    # タイトル行を挿入
    ws_sum.insert_rows(1)
    title_cell = ws_sum.cell(row=1, column=1, value=f"就業規則サマリー　{company}　{datetime.now().strftime('%Y年%m月%d日')}")
    title_cell.font = Font(bold=True, size=13, color="0D1B30")
    ws_sum.merge_cells(start_row=1, start_column=1, end_row=1, end_column=5)
    ws_sum.row_dimensions[1].height = 30

    # データ行
    for ri, (item, _) in enumerate(SEISHAIN_SUMMARY_ITEMS, 3):
        data = summary.get(item, {})
        if isinstance(data, dict):
            jibun   = data.get("条文", "—")
            naiyou  = data.get("内容", "記載なし")
            yocheck = data.get("要確認", False)
        else:
            jibun, naiyou, yocheck = "—", str(data), False

        fill = red_fill if yocheck or naiyou == "記載なし" else gold_fill
        row_data = [item, jibun, naiyou, "⚠️" if yocheck else "", ""]
        for ci, val in enumerate(row_data, 1):
            cell = ws_sum.cell(row=ri, column=ci, value=val)
            cell.fill = fill
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            if ci == 1:
                cell.font = Font(bold=True, size=10)
            else:
                cell.font = Font(size=10)
        ws_sum.row_dimensions[ri].height = 40

    buf_out = io.BytesIO()
    wb.save(buf_out)
    return buf_out.getvalue()


# ============================================================
# Word生成
# ============================================================

def build_docx(report_text: str, company: str, label: str, phase: str) -> bytes:
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return b""

    items = parse_report(report_text)
    doc = DocxDocument()

    # ページ設定（A4・余白2cm）XMLを直接書き換え
    from lxml import etree as _etree
    body = doc.element.body
    sectPr = body.get_or_add_sectPr()
    pgSz = sectPr.find(qn("w:pgSz"))
    if pgSz is None:
        pgSz = OxmlElement("w:pgSz")
        sectPr.insert(0, pgSz)
    pgSz.set(qn("w:w"), "11906")   # A4幅
    pgSz.set(qn("w:h"), "16838")   # A4高
    pgMar = sectPr.find(qn("w:pgMar"))
    if pgMar is None:
        pgMar = OxmlElement("w:pgMar")
        sectPr.insert(1, pgMar)
    pgMar.set(qn("w:top"),    "1134")   # 2cm
    pgMar.set(qn("w:right"),  "1134")
    pgMar.set(qn("w:bottom"), "1134")
    pgMar.set(qn("w:left"),   "1134")
    pgMar.set(qn("w:header"), "708")
    pgMar.set(qn("w:footer"), "708")
    pgMar.set(qn("w:gutter"), "0")

    # タイトル
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("就業規則チェックレポート")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x30)

    # メタ情報
    doc.add_paragraph()
    meta_items = [
        ("会社名", company),
        ("対象コース", label),
        ("フェーズ", PHASE_LABELS.get(phase, phase)),
        ("出力日時", datetime.now().strftime("%Y年%m月%d日 %H:%M")),
    ]
    for k, v in meta_items:
        p = doc.add_paragraph()
        run_k = p.add_run(f"{k}：")
        run_k.bold = True
        run_k.font.size = Pt(10)
        run_v = p.add_run(v)
        run_v.font.size = Pt(10)

    doc.add_paragraph()

    # サマリー
    h = doc.add_paragraph("チェックサマリー")
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(13)
    h.runs[0].font.color.rgb = RGBColor(0x0D, 0x1B, 0x30)

    counts = {k: 0 for k in ALERT_CONFIG}
    for it in items:
        if it["level"] in counts:
            counts[it["level"]] += 1
    for key, cfg in ALERT_CONFIG.items():
        p = doc.add_paragraph()
        p.add_run(f"  {cfg['icon']} {cfg['label']}：{counts[key]}件")
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # 詳細
    h2 = doc.add_paragraph("詳細チェック結果")
    h2.runs[0].bold = True
    h2.runs[0].font.size = Pt(13)
    h2.runs[0].font.color.rgb = RGBColor(0x0D, 0x1B, 0x30)

    level_rgb = {
        "CRITICAL":    RGBColor(0xC5, 0x3B, 0x3B),
        "WARNING":     RGBColor(0xB7, 0x79, 0x10),
        "CAUTION":     RGBColor(0xC0, 0x5A, 0x10),
        "HUMAN_CHECK": RGBColor(0x3B, 0x5B, 0xC5),
        "OK":          RGBColor(0x27, 0x7A, 0x4A),
    }

    for it in items:
        cfg = ALERT_CONFIG.get(it["level"], {"icon": "ℹ️", "label": it["level"]})
        rgb = level_rgb.get(it["level"], RGBColor(0x33, 0x33, 0x33))

        # 項目タイトル
        p = doc.add_paragraph()
        run = p.add_run(f'{cfg["icon"]} {it["title"]}')
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x30)

        # 判定バッジ
        p2 = doc.add_paragraph()
        badge = p2.add_run(f'　判定：{cfg["icon"]} {cfg["label"]}')
        badge.bold = True
        badge.font.size = Pt(10)
        badge.font.color.rgb = rgb

        # 各フィールド
        fields = [
            ("❗ 問題の内容", it["問題の内容"]),
            ("📌 該当箇所",   it["該当箇所"]),
            ("✏️ 修正案",    it["修正案"]),
            ("👁️ 審査官の目線", it["審査官の目線"]),
            ("📎 根拠",       it["根拠"]),
        ]
        for fname, fval in fields:
            if fval:
                # フィールドラベル
                p_label = doc.add_paragraph()
                label_run = p_label.add_run(fname)
                label_run.bold = True
                label_run.font.size = Pt(9.5)
                label_run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x30)
                p_label.paragraph_format.space_after = Pt(2)
                # フィールド内容（改行を段落に変換）
                for line in fval.split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    p_val = doc.add_paragraph()
                    p_val.paragraph_format.left_indent = Pt(16)
                    p_val.paragraph_format.space_after = Pt(1)
                    val_run = p_val.add_run(line)
                    val_run.font.size = Pt(9.5)

        # 区切り線代わりの空行
        sep = doc.add_paragraph()
        sep.paragraph_format.space_after = Pt(4)
        pPr = sep._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "DDDDDD")
        pBdr.append(bottom)
        pPr.append(pBdr)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_docx_with_summary(report_text: str, company: str, label: str,
                              phase: str, summary: dict) -> bytes:
    """サマリーページ付きWord生成（正社員化コース専用）"""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re
    except ImportError:
        return build_docx(report_text, company, label, phase)

    doc = DocxDocument()

    # ページ設定（XMLで直接指定）
    body = doc.element.body
    sectPr = body.get_or_add_sectPr()
    pgSz = sectPr.find(qn("w:pgSz"))
    if pgSz is None:
        pgSz = OxmlElement("w:pgSz")
        sectPr.insert(0, pgSz)
    pgSz.set(qn("w:w"), "11906")
    pgSz.set(qn("w:h"), "16838")
    pgMar = sectPr.find(qn("w:pgMar"))
    if pgMar is None:
        pgMar = OxmlElement("w:pgMar")
        sectPr.insert(1, pgMar)
    for attr, val in [("w:top","1134"),("w:right","1134"),("w:bottom","1134"),
                       ("w:left","1134"),("w:header","708"),("w:footer","708"),("w:gutter","0")]:
        pgMar.set(qn(attr), val)

    def add_shading(cell, fill_hex):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_hex)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    # ── サマリーページ ──
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("就業規則サマリー（正社員）")
    r.bold = True; r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x0D, 0x1B, 0x30)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run(f"{company}　{datetime.now().strftime('%Y年%m月%d日')}").font.size = Pt(10)

    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run("⚠️ 赤背景の項目は「記載なし」または「要確認」です。労働条件通知書・出勤簿・給与明細と照合してください。")
    nr.font.size = Pt(9); nr.font.color.rgb = RGBColor(0xC5, 0x3B, 0x3B)
    doc.add_paragraph()

    # サマリーテーブル
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for ci, txt in enumerate(["項目", "条文番号", "就業規則の記載内容", "要確認"]):
        hdr[ci].text = txt
        add_shading(hdr[ci], "0D1B30")
        if hdr[ci].paragraphs[0].runs:
            run = hdr[ci].paragraphs[0].runs[0]
            run.bold = True; run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC8, 0xA8, 0x4A)

    for item, _ in SEISHAIN_SUMMARY_ITEMS:
        data = summary.get(item, {})
        if isinstance(data, dict):
            jibun   = data.get("条文", "—")
            naiyou  = data.get("内容", "記載なし")
            yocheck = data.get("要確認", False)
        else:
            jibun, naiyou, yocheck = "—", str(data), False

        row = tbl.add_row().cells
        row[0].text = item; row[1].text = jibun
        row[2].text = naiyou; row[3].text = "⚠️" if yocheck else ""
        fill = "FFE4E4" if (yocheck or naiyou == "記載なし") else "FFFFFF"
        for cell in row:
            add_shading(cell, fill)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
        if row[0].paragraphs[0].runs:
            row[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ── チェック結果ページ ──
    at = doc.add_paragraph()
    at.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = at.add_run("就業規則チェックレポート")
    r2.bold = True; r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x0D, 0x1B, 0x30)

    level_rgb = {
        "CRITICAL": RGBColor(0xC5,0x3B,0x3B), "WARNING": RGBColor(0xB7,0x79,0x10),
        "CAUTION":  RGBColor(0xC0,0x5A,0x10), "HUMAN_CHECK": RGBColor(0x3B,0x5B,0xC5),
        "OK":       RGBColor(0x27,0x7A,0x4A),
    }
    icons = {"CRITICAL":"🔴","WARNING":"🟡","CAUTION":"🟠","HUMAN_CHECK":"👤","OK":"🟢"}

    for sec in re.split(r"\n#{2,4} ", report_text)[1:]:
        lines_s = sec.strip().split("\n")
        sec_title = lines_s[0].strip()
        sec_body  = "\n".join(lines_s[1:])
        level = None
        for key in ["CRITICAL","WARNING","CAUTION","HUMAN_CHECK","OK"]:
            if re.search(rf"判定[：:][^\n]{{0,40}}{key}", sec_body):
                level = key; break
        if level is None:
            continue
        icon = icons.get(level,""); rgb = level_rgb.get(level, RGBColor(0x33,0x33,0x33))
        p = doc.add_paragraph()
        rt = p.add_run(f"{icon} {sec_title}")
        rt.bold = True; rt.font.size = Pt(10.5)
        rt.font.color.rgb = RGBColor(0x0D,0x1B,0x30)
        p2 = doc.add_paragraph()
        rl = p2.add_run(f"　判定：{icon} {level}")
        rl.bold = True; rl.font.size = Pt(9.5); rl.font.color.rgb = rgb

        for fname in ["問題の内容","該当箇所","修正案","根拠"]:
            esc = re.escape(fname)
            m = re.search(rf"\*\*{esc}[：:]\*\*[ \t]*\n?(.*?)(?=\n\*\*[^\n]+[：:]\*\*|\n---|\Z)", sec_body, re.DOTALL)
            val = m.group(1).strip() if m else ""
            if not val:
                continue
            pl = doc.add_paragraph()
            pl.paragraph_format.left_indent = Pt(8)
            rl2 = pl.add_run(f"【{fname}】")
            rl2.bold = True; rl2.font.size = Pt(9)
            for line in val.split("\n"):
                if not line.strip(): continue
                pv = doc.add_paragraph()
                pv.paragraph_format.left_indent = Pt(20)
                pv.paragraph_format.space_after = Pt(1)
                pv.add_run(line.strip()).font.size = Pt(9)
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
# サイドバー
# ============================================================

with st.sidebar:
    st.markdown(
        '<div class="sb-brand">'
        '<div class="sb-brand-eye">HUMAX | AUDIT SYSTEM</div>'
        '<div class="sb-brand-title">就業規則チェックツール</div>'
        '<div class="sb-brand-sub">社会保険労務士法人ヒューマックス</div>'
        '</div>',
        unsafe_allow_html=True,
    )

    available_keys = get_available_keys()

    st.markdown("**利用可能な判定ルール**")
    if available_keys:
        for year, date, cid in sorted(available_keys):
            c = next((x for x in COURSES if x["id"] == cid), None)
            label = f"{year}年度 {date[:2]}月{date[2:]}日〜　{c['name'] if c else cid}"
            st.markdown(
                f'<span class="sb-rule-item">&#10003;&#160;{label}</span>',
                unsafe_allow_html=True,
            )
    else:
        st.warning("判定ルールJSONがありません")

    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("**チェックフェーズ**")
    for pk, (pnum, pshort, pdesc) in PHASE_ITEMS.items():
        st.markdown(f"**{pnum}　{pshort}**")
        st.caption(pdesc)

    st.markdown(
        '<div class="sb-footer">AuditX v0.5　Powered by Gemini API</div>',
        unsafe_allow_html=True,
    )


# ============================================================
# メイン：ページヘッダー
# ============================================================

st.markdown(
    '<div class="page-eyebrow">HUMAX | LABOR REGULATION AUDIT SYSTEM</div>'
    '<div class="page-title">就業規則チェックツール（Humax）</div>'
    '<div class="page-subtitle">'
    '社会保険労務士法人ヒューマックス向け就業規則チェックシステム'
    '　— 業務効率の最大化のために'
    '</div>',
    unsafe_allow_html=True,
)
st.markdown('<hr class="page-divider">', unsafe_allow_html=True)


# ============================================================
# STEP 1 : 会社情報
# ============================================================

st.markdown('<div class="sec-label">STEP 1　会社情報</div>', unsafe_allow_html=True)

col1, col2 = st.columns([3, 2])
with col1:
    company_name = st.text_input(
        "会社名", placeholder="例：株式会社〇〇", label_visibility="collapsed"
    )
with col2:
    industry = st.text_input(
        "業種（任意）", placeholder="例：小売業、サービス業", label_visibility="collapsed"
    )

st.markdown('<hr class="divider">', unsafe_allow_html=True)


# ============================================================
# STEP 2 : フェーズ選択
# ============================================================

st.markdown('<div class="sec-label">STEP 2　チェックフェーズ</div>', unsafe_allow_html=True)

if "selected_phase" not in st.session_state:
    st.session_state.selected_phase = "phase1"

cols_p = st.columns(3)
for col, (pk, (pnum, pshort, pdesc)) in zip(cols_p, PHASE_ITEMS.items()):
    with col:
        is_active = st.session_state.selected_phase == pk
        card_cls  = "phase-active" if is_active else "phase-inactive"
        num_cls   = "phase-num-active" if is_active else "phase-num-inactive"
        ttl_cls   = "phase-title-active" if is_active else "phase-title-inactive"
        dsc_cls   = "phase-desc-active" if is_active else "phase-desc-inactive"

        st.markdown(
            f'<div class="{card_cls}">'
            f'<div class="{num_cls}">{pnum}</div>'
            f'<div class="{ttl_cls}">{pshort}</div>'
            f'<div class="{dsc_cls}">{pdesc}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

        btn_wrap = "btn-active" if is_active else "btn-inactive"
        st.markdown(f'<div class="{btn_wrap}">', unsafe_allow_html=True)
        if st.button(
            "選択中 ✓" if is_active else "選択する",
            key=f"pb_{pk}",
            use_container_width=True,
        ):
            st.session_state.selected_phase = pk
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

selected_phase = st.session_state.selected_phase
_, pshort, pdesc = PHASE_ITEMS[selected_phase]
st.markdown(
    f'<div class="phase-hint">&#128203;&#160;<strong>{pshort}</strong>　{pdesc}</div>',
    unsafe_allow_html=True,
)

st.markdown('<hr class="divider">', unsafe_allow_html=True)


# ============================================================
# STEP 3 : コース選択（コースごとに年度・施行日）
# ============================================================

st.markdown(
    '<div class="sec-label">STEP 3　申請予定の助成金コース（コースごとに年度・施行日を設定）</div>',
    unsafe_allow_html=True,
)

selected_courses = []
prev_group = None

for course in COURSES:
    if course["group"] != prev_group:
        st.markdown(
            f'<div class="group-label">{course["group"]}</div>',
            unsafe_allow_html=True,
        )
        prev_group = course["group"]

    checked = st.checkbox(course["name"], key=f"chk_{course['id']}")

    if checked:
        c1, c2, c3 = st.columns([2, 2, 3])
        with c1:
            year = st.selectbox(
                "年度",
                YEAR_OPTIONS,
                format_func=lambda x: f"令和{x[1:]}年度",
                key=f"y_{course['id']}",
            )
        with c2:
            date_opts = DATE_OPTIONS.get(year, ["0401"])
            date = st.selectbox(
                "施行日",
                date_opts,
                format_func=lambda x: f"{x[:2]}月{x[2:]}日以降",
                key=f"d_{course['id']}",
            )
        with c3:
            rule_exists = (year, date, course["id"]) in available_keys
            if rule_exists:
                st.success("判定ルール：あり ✓")
            else:
                st.warning(
                    f"令和{year[1:]}年度（{date[:2]}月{date[2:]}日〜）の"
                    "判定ルールJSONがありません"
                )

        selected_courses.append({
            "course_id":   course["id"],
            "name":        course["name"],
            "group":       course["group"],
            "year":        year,
            "date":        date,
            "rule_exists": rule_exists,
        })

st.markdown('<hr class="divider">', unsafe_allow_html=True)


# ============================================================
# STEP 3.5 : 正社員化コース対象者情報
# ============================================================

ca_seishain_selected = any(c["course_id"] == "CA_seishain" for c in selected_courses)

target_person_info = {}

if ca_seishain_selected:
    st.markdown(
        '<div class="sec-label">STEP 3.5　正社員化コース：対象者情報（必須）</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        '<div class="phase-hint">&#9888;&#160;' 
        '<strong>正社員化コースは対象者の情報が審査に直結します。</strong>'
        '　複数名いる場合は、最も不利な条件（最年長・最古入社・最も近い転換予定日）の方の情報を入力してください。'
        '</div>',
        unsafe_allow_html=True,
    )

    tp_col1, tp_col2, tp_col3 = st.columns(3)
    # 和暦変換テーブル（全日付入力で共通使用）
    WAREKI = [
        ("令和", 2019, "R"),
        ("平成", 1989, "H"),
        ("昭和", 1926, "S"),
        ("大正", 1912, "T"),
    ]

    with tp_col1:
        st.markdown("**対象者の生年月日**（最年長者）")

        dob_mode = st.radio(
            "入力方式", ["西暦", "和暦"],
            horizontal=True, key="dob_mode", label_visibility="collapsed"
        )

        target_dob = None
        if dob_mode == "西暦":
            target_dob = st.date_input(
                "生年月日（西暦）",
                value=None,
                min_value=datetime(1940, 1, 1).date(),
                max_value=datetime.now().date(),
                label_visibility="collapsed",
                key="target_dob",
            )
        else:
            w_col1, w_col2, w_col3, w_col4 = st.columns([2, 1, 1, 1])
            with w_col1:
                gengou = st.selectbox("元号", [g for g, _, _ in WAREKI],
                                      key="dob_gengou", label_visibility="collapsed")
            with w_col2:
                wareki_year = st.number_input("年", min_value=1, max_value=99,
                                              value=1, key="dob_wy", label_visibility="collapsed")
            with w_col3:
                dob_month = st.number_input("月", min_value=1, max_value=12,
                                            value=1, key="dob_m", label_visibility="collapsed")
            with w_col4:
                dob_day = st.number_input("日", min_value=1, max_value=31,
                                          value=1, key="dob_d", label_visibility="collapsed")
            # 西暦に変換
            base_year = next(base for g, base, _ in WAREKI if g == gengou)
            seireki_year = base_year + wareki_year - 1
            try:
                target_dob = datetime(seireki_year, dob_month, dob_day).date()
                st.caption(f"→ 西暦 {seireki_year}年{dob_month}月{dob_day}日")
            except ValueError:
                st.error("日付が正しくありません")
                target_dob = None
    with tp_col2:
        st.markdown("**対象者の入社日**（最も古い日付）")
        hire_mode = st.radio(
            "入力方式", ["西暦", "和暦"],
            horizontal=True, key="hire_mode", label_visibility="collapsed"
        )
        target_hire = None
        if hire_mode == "西暦":
            target_hire = st.date_input(
                "入社日（西暦）",
                value=None,
                min_value=datetime(1980, 1, 1).date(),
                max_value=datetime.now().date(),
                label_visibility="collapsed",
                key="target_hire",
            )
        else:
            h_col1, h_col2, h_col3, h_col4 = st.columns([2, 1, 1, 1])
            with h_col1:
                hire_gengou = st.selectbox("元号", [g for g, _, _ in WAREKI],
                                           key="hire_gengou", label_visibility="collapsed")
            with h_col2:
                hire_wy = st.number_input("年", min_value=1, max_value=99,
                                          value=1, key="hire_wy", label_visibility="collapsed")
            with h_col3:
                hire_m = st.number_input("月", min_value=1, max_value=12,
                                         value=1, key="hire_m", label_visibility="collapsed")
            with h_col4:
                hire_d = st.number_input("日", min_value=1, max_value=31,
                                         value=1, key="hire_d", label_visibility="collapsed")
            hire_base = next(base for g, base, _ in WAREKI if g == hire_gengou)
            hire_year = hire_base + hire_wy - 1
            try:
                target_hire = datetime(hire_year, hire_m, hire_d).date()
                st.caption(f"→ 西暦 {hire_year}年{hire_m}月{hire_d}日")
            except ValueError:
                st.error("日付が正しくありません")
                target_hire = None

    with tp_col3:
        st.markdown("**転換予定日**")
        convert_mode = st.radio(
            "入力方式", ["西暦", "和暦"],
            horizontal=True, key="convert_mode", label_visibility="collapsed"
        )
        target_convert = None
        if convert_mode == "西暦":
            target_convert = st.date_input(
                "転換予定日（西暦）",
                value=None,
                min_value=datetime(2020, 1, 1).date(),
                max_value=datetime(2035, 12, 31).date(),
                label_visibility="collapsed",
                key="target_convert",
            )
        else:
            c_col1, c_col2, c_col3, c_col4 = st.columns([2, 1, 1, 1])
            with c_col1:
                conv_gengou = st.selectbox("元号", ["令和", "平成"],
                                           key="conv_gengou", label_visibility="collapsed")
            with c_col2:
                conv_wy = st.number_input("年", min_value=1, max_value=30,
                                          value=7, key="conv_wy", label_visibility="collapsed")
            with c_col3:
                conv_m = st.number_input("月", min_value=1, max_value=12,
                                         value=1, key="conv_m", label_visibility="collapsed")
            with c_col4:
                conv_d = st.number_input("日", min_value=1, max_value=31,
                                         value=1, key="conv_d", label_visibility="collapsed")
            conv_base = 2019 if conv_gengou == "令和" else 1989
            conv_year = conv_base + conv_wy - 1
            try:
                target_convert = datetime(conv_year, conv_m, conv_d).date()
                st.caption(f"→ 西暦 {conv_year}年{conv_m}月{conv_d}日")
            except ValueError:
                st.error("日付が正しくありません")
                target_convert = None

    # 入力値の計算
    if target_dob and target_hire and target_convert:
        today = datetime.now().date()
        age = today.year - target_dob.year - (
            (today.month, today.day) < (target_dob.month, target_dob.day)
        )
        hire_months = (today.year - target_hire.year) * 12 + (today.month - target_hire.month)
        convert_days = (target_convert - today).days

        info_cols = st.columns(3)
        with info_cols[0]:
            st.metric("現在の年齢", f"{age}歳")
        with info_cols[1]:
            st.metric("在籍期間", f"{hire_months}ヶ月（{hire_months//12}年{hire_months%12}ヶ月）")
        with info_cols[2]:
            st.metric("転換まで", f"{convert_days}日")

        # 簡易要件チェック
        warnings = []
        if hire_months < 6:
            warnings.append("⚠️ 在籍期間が6ヶ月未満です（正社員化コースは原則6ヶ月以上の雇用が必要）")
        if convert_days < 0:
            warnings.append("⚠️ 転換予定日が過去の日付になっています")
        for w in warnings:
            st.warning(w)

        target_person_info = {
            "生年月日": target_dob.strftime("%Y年%m月%d日"),
            "年齢": f"{age}歳",
            "入社日": target_hire.strftime("%Y年%m月%d日"),
            "在籍期間": f"{hire_months}ヶ月",
            "転換予定日": target_convert.strftime("%Y年%m月%d日"),
            "転換まで": f"{convert_days}日",
        }
    else:
        st.info("対象者の生年月日・入社日・転換予定日をすべて入力してください。")
        target_person_info = {}

    st.markdown('<hr class="divider">', unsafe_allow_html=True)


# ============================================================
# STEP 4 : ファイルアップロード
# ============================================================

st.markdown(
    '<div class="sec-label">STEP 4　就業規則ファイルをアップロード</div>',
    unsafe_allow_html=True,
)

if selected_phase == "phase3":
    st.info(
        "Phase 3 では導入前・導入後・その他改訂分すべての就業規則を"
        "まとめてアップロードしてください。"
    )

col_u1, col_u2 = st.columns(2)
with col_u1:
    st.markdown("**就業規則**（必須・複数可）")
    st.caption("就業規則・賃金規程・育児介護休業規程 等")
    uploaded_files = st.file_uploader(
        "就業規則",
        type=["pdf", "docx", "doc"],
        accept_multiple_files=True,
        label_visibility="collapsed",
        key="main_files",
    )
with col_u2:
    st.markdown("**補足書類**（任意）")
    st.caption("賃金台帳・労働条件通知書・労使協定書 等")
    supplementary_files = st.file_uploader(
        "補足書類",
        type=["pdf", "docx", "doc", "txt"],
        accept_multiple_files=True,
        label_visibility="collapsed",
        key="sup_files",
    )

if uploaded_files:
    for f in uploaded_files:
        st.caption(f"&#128196;&#160;{f.name}　({f.size/1024:.1f} KB)")

with st.expander("担当者メモ（任意）"):
    st.text_area(
        "メモ",
        placeholder="例：転換予定者は〇〇さん（入社3年目）。定年規定の確認が特に重要。",
        height=70,
        label_visibility="collapsed",
    )

st.markdown('<hr class="divider">', unsafe_allow_html=True)


# ============================================================
# 実行ボタン
# ============================================================

valid_courses = [c for c in selected_courses if c["rule_exists"]]
# 正社員化コース選択時は対象者情報も必須
ca_selected_in_valid = any(c["course_id"] == "CA_seishain" for c in valid_courses)
person_info_ok = (not ca_selected_in_valid) or bool(target_person_info)
can_run = bool(company_name) and bool(valid_courses) and bool(uploaded_files) and person_info_ok

missing = []
if not company_name:   missing.append("会社名")
if not valid_courses:  missing.append("助成金コース")
if not uploaded_files: missing.append("就業規則ファイル")
if ca_selected_in_valid and not person_info_ok: missing.append("対象者情報（生年月日・入社日・転換予定日）")
if missing:
    st.caption("未入力の項目：" + "　/　".join(missing))

run_button = st.button(
    "チェック開始",
    disabled=not can_run,
    type="primary",
    use_container_width=True,
)

if not can_run:
    st.caption("就業規則ファイルをアップロードすると「チェック開始」が有効になります。")


# ============================================================
# チェック実行
# ============================================================

# session_stateにレポートを保持（ダウンロード後のリセット防止）
if "saved_reports" not in st.session_state:
    st.session_state.saved_reports = []
if "saved_meta" not in st.session_state:
    st.session_state.saved_meta = {}

if run_button and can_run:

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<div class="sec-label">チェック結果</div>', unsafe_allow_html=True)

    # APIキー
    api_key = st.secrets.get("GOOGLE_API_KEY", os.getenv("GOOGLE_API_KEY"))
    if not api_key:
        env_path = Path(__file__).parent / ".env"
        if env_path.exists():
            with open(env_path, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if line.startswith("GOOGLE_API_KEY=") and "ここに" not in line:
                        api_key = line.split("=", 1)[1].strip()
                        break
    if not api_key:
        st.error("GOOGLE_API_KEY が設定されていません。Streamlit Cloud の Secrets を確認してください。")
        st.stop()
    os.environ["GOOGLE_API_KEY"] = api_key

    # テキスト抽出
    all_texts, target_filenames = [], []
    all_upload = list(uploaded_files) + (list(supplementary_files) if supplementary_files else [])

    with st.status("ファイルを読み込み中...", expanded=True) as status:
        for uf in all_upload:
            st.write(f"読み込み中：{uf.name}")
            try:
                suffix = Path(uf.name).suffix
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    tmp.write(uf.read())
                    tmp_path = tmp.name
                text = extract_text(tmp_path)
                os.unlink(tmp_path)
                all_texts.append(f"=== {uf.name} ===\n{text}")
                target_filenames.append(uf.name)
                st.write(f"完了：{uf.name}（{len(text):,} 文字）")
            except Exception as e:
                st.error(f"エラー：{uf.name}　{e}")
        status.update(label="読み込み完了", state="complete")

    if not all_texts:
        st.stop()

    combined_text = "\n\n".join(all_texts)
    all_reports = []

    for ci in valid_courses:
        with st.status(f"判定中：{ci['name']}", expanded=True) as status:
            rk = load_rule_knowledge(ci["course_id"], ci["year"], ci["date"])
            if not rk:
                st.error("判定ルールJSONの読み込みに失敗しました")
                status.update(label=f"エラー：{ci['name']}", state="error")
                continue
            st.write("判定ルール読み込み完了")
            st.write("Gemini API で判定中（1〜2分ほどかかります）...")
            try:
                others = [
                    f"{c['group']}　{c['name']}"
                    for c in valid_courses if c["course_id"] != ci["course_id"]
                ]
                # 正社員化コースの場合は対象者情報を渡す
                person_info = target_person_info if ci["course_id"] == "CA_seishain" else {}
                result = run_audit(
                    combined_text, rk, ci["course_id"],
                    selected_phase, others or None,
                    target_person_info=person_info,
                )
                report = generate_report(
                    result, company_name, ci["course_id"],
                    selected_phase, rk, target_filenames,
                )
                # 正社員化コースはサマリーも抽出
                seishain_summary = {}
                if ci["course_id"] == "CA_seishain":
                    st.write("就業規則サマリーを抽出中...")
                    try:
                        seishain_summary = extract_seishain_summary(combined_text, api_key)
                        if seishain_summary:
                            st.write(f"サマリー抽出完了（{len(seishain_summary)}項目）")
                        else:
                            st.warning("サマリーの抽出結果が空でした。再度お試しください。")
                    except Exception as e_sum:
                        st.error(f"サマリー抽出エラー：{e_sum}")

                all_reports.append({
                    "course_id": ci["course_id"],
                    "label":     f"{ci['group']}　{ci['name']}",
                    "report":    report,
                    "seishain_summary": seishain_summary,
                })
                status.update(label=f"完了：{ci['name']}", state="complete")
            except Exception as e:
                st.error(f"APIエラー：{e}")
                status.update(label=f"エラー：{ci['name']}", state="error")

    if all_reports:
        st.session_state.saved_reports = all_reports
        st.session_state.saved_meta = {
            "company": company_name,
            "phase": selected_phase,
            "now_str": datetime.now().strftime("%Y%m%d_%H%M"),
        }
        st.success(f"チェック完了　{len(all_reports)} コース")
        now_str = datetime.now().strftime("%Y%m%d_%H%M")

        report_tabs = st.tabs([r["label"] for r in all_reports]) if len(all_reports) > 1 else [st.container()]
        for tab, rd in zip(report_tabs, all_reports):
            with tab:
                # ── レポートをパースして構造化表示 ──
                display_report_visual(rd["report"], rd["label"])

                st.markdown("---")
                st.markdown("**ダウンロード**")
                dl_cols = st.columns(3)

                # Excel
                with dl_cols[0]:
                    _summary = rd.get("seishain_summary", {})
                    if rd["course_id"] == "CA_seishain" and _summary:
                        xlsx_data = build_xlsx_with_summary(rd["report"], company_name, rd["label"], selected_phase, _summary)
                    else:
                        xlsx_data = build_xlsx(rd["report"], company_name, rd["label"], selected_phase)
                    st.download_button(
                        "📊 Excel (.xlsx)",
                        data=xlsx_data,
                        file_name=f"AuditX_{company_name}_{rd['course_id']}_{selected_phase}_{now_str}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                        key=f"xl_{rd['course_id']}",
                    )

                # Word
                with dl_cols[1]:
                    _summary = rd.get("seishain_summary", {})
                    if rd["course_id"] == "CA_seishain" and _summary:
                        docx_data = build_docx_with_summary(rd["report"], company_name, rd["label"], selected_phase, _summary)
                    else:
                        docx_data = build_docx(rd["report"], company_name, rd["label"], selected_phase)
                    st.download_button(
                        "📝 Word (.docx)",
                        data=docx_data,
                        file_name=f"AuditX_{company_name}_{rd['course_id']}_{selected_phase}_{now_str}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key=f"dx_{rd['course_id']}",
                    )

                # Markdown
                with dl_cols[2]:
                    st.download_button(
                        "📄 Markdown (.md)",
                        data=rd["report"].encode("utf-8"),
                        file_name=f"AuditX_{company_name}_{rd['course_id']}_{selected_phase}_{now_str}.md",
                        mime="text/markdown",
                        use_container_width=True,
                        key=f"md_{rd['course_id']}",
                    )

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.caption("就業規則チェックツール（ヒューマックス）　v0.5")
    st.caption(
        "※ 本ツールは実務補助用です。最終判断は必ず担当社会保険労務士が行ってください。"
    )

# ── チェック済みレポートが session_state にある場合は常に表示 ──
elif st.session_state.get("saved_reports"):
    all_reports = st.session_state.saved_reports
    meta = st.session_state.saved_meta
    company_name_s = meta.get("company", "")
    selected_phase_s = meta.get("phase", "phase1")
    now_str = meta.get("now_str", datetime.now().strftime("%Y%m%d_%H%M"))

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<div class="sec-label">チェック結果（保存済み）</div>', unsafe_allow_html=True)
    st.info("💾 前回のチェック結果が保存されています。新しいチェックを行うと上書きされます。")

    report_tabs = st.tabs([r["label"] for r in all_reports]) if len(all_reports) > 1 else [st.container()]
    for tab, rd in zip(report_tabs, all_reports):
        with tab:
            display_report_visual(rd["report"], rd["label"])
            st.markdown("---")
            st.markdown("**ダウンロード**")
            dl_cols = st.columns(3)
            with dl_cols[0]:
                _sum2 = rd.get("seishain_summary", {})
                if rd["course_id"] == "CA_seishain" and _sum2:
                    xlsx_data = build_xlsx_with_summary(rd["report"], company_name_s, rd["label"], selected_phase_s, _sum2)
                else:
                    xlsx_data = build_xlsx(rd["report"], company_name_s, rd["label"], selected_phase_s)
                st.download_button("📊 Excel (.xlsx)", data=xlsx_data,
                    file_name=f"AuditX_{company_name_s}_{rd['course_id']}_{selected_phase_s}_{now_str}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True, key=f"xl2_{rd['course_id']}")
            with dl_cols[1]:
                _sum2 = rd.get("seishain_summary", {})
                if rd["course_id"] == "CA_seishain" and _sum2:
                    docx_data = build_docx_with_summary(rd["report"], company_name_s, rd["label"], selected_phase_s, _sum2)
                else:
                    docx_data = build_docx(rd["report"], company_name_s, rd["label"], selected_phase_s)
                st.download_button("📝 Word (.docx)", data=docx_data,
                    file_name=f"AuditX_{company_name_s}_{rd['course_id']}_{selected_phase_s}_{now_str}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True, key=f"dx2_{rd['course_id']}")
            with dl_cols[2]:
                st.download_button("📄 Markdown (.md)", data=rd["report"].encode("utf-8"),
                    file_name=f"AuditX_{company_name_s}_{rd['course_id']}_{selected_phase_s}_{now_str}.md",
                    mime="text/markdown", use_container_width=True, key=f"md2_{rd['course_id']}")

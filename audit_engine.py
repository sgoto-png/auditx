"""
audit_engine.py - 就業規則チェックエンジン（Gemini API版）
"""

import argparse
import json
import os
import sys
from datetime import datetime
from pathlib import Path

from google import genai
from google.genai import types
import fitz  # PyMuPDF

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ============================================================
# 定数
# ============================================================

MODEL = "gemini-2.5-flash"
MAX_TOKENS = 65536

PHASE_LABELS = {
    "phase1": "Phase1：助成金申請準備開始時（就業規則新規作成後）",
    "phase2": "Phase2：制度導入・規定改訂時",
    "phase3": "Phase3：支給申請提出前（全バージョン整合性確認）",
}

ALERT_ICONS = {
    "CRITICAL": "🔴",
    "WARNING": "🟡",
    "CAUTION": "🟠",
    "HUMAN_CHECK": "👤",
    "OK": "🟢",
    "INFO": "ℹ️",
}

# ============================================================
# テキスト抽出
# ============================================================

def extract_text_from_pdf(pdf_path: str) -> str:
    doc = fitz.open(pdf_path)
    text_parts = []
    for page_num, page in enumerate(doc, 1):
        text = page.get_text()
        if text.strip():
            text_parts.append(f"--- {page_num}ページ ---\n{text}")
    doc.close()
    return "\n".join(text_parts)

def extract_text_from_docx(docx_path: str) -> str:
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx がインストールされていません。")
    document = docx.Document(docx_path)
    text_parts = []
    for para in document.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts)

def extract_text(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"対応していないファイル形式です: {ext}")

# ============================================================
# プロンプト生成
# ============================================================

def build_prompt(rule_knowledge: dict, course_id: str, phase: str,
                 rules_text: str, additional_courses: list = None,
                 target_person_info: dict = None,
                 humax_knowledge: dict = None,
                 ca_shoyo_info: dict = None,
                 law_knowledge: dict = None) -> str:
    course_name = rule_knowledge.get("meta", {}).get("course_name", course_id)
    year = rule_knowledge.get("meta", {}).get("year", "")
    eff_date = rule_knowledge.get("meta", {}).get("effective_date", "")
    rule_json = json.dumps(rule_knowledge, ensure_ascii=False, indent=2)

    phase_instruction = {
        "phase1": """
【Phase1 チェック内容】
A. 全コース共通チェック
1. 従業員の呼称が全規程で統一されているか
2. 雇用形態の定義が全規程で矛盾がないか
3. 全従業員に適用される規程があるか
4. 支給されている手当が全て定義されているか
5. 手当の定義（残業代算定含否・固定/変動の別）が明確か
6. 雇用形態別の給与形態規定に誤りがないか
7. 最新の労働基準法に抵触する内容がないか
8. 誤字脱字

B. 助成金コース別チェック
- 申請予定の助成金の要件を満たすための規定が存在するか
- 定年規定が助成金申請の妨げになっていないか
- 昇給・賞与・退職金の規定が助成金要件と矛盾しないか
""",
        "phase2": """
【Phase2 チェック内容】
A. 全コース共通チェック（Phase1と同様）
B. 導入・改訂内容のチェック
- 今回新たに導入・改訂した規定の内容が支給要領の要件を満たしているか
- 必須記載事項が全て含まれているか
- NG記載パターンが含まれていないか
C. タイムライン確認
- 制度導入のタイミングが支給要領の要件を満たしているか
""",
        "phase3": """
【Phase3 チェック内容】
A. 全コース共通チェック（Phase1と同様）
B. 全バージョン間の整合性チェック
- 導入前規則が「取組前要件」を満たしているか
- 導入後規則が「取組後要件」を満たしているか
- バージョン間で意図しない矛盾・変更がないか
C. 支給申請書類との整合性
- 賃金台帳・労働条件通知書との照合が必要な項目を明示
""",
    }.get(phase, "")

    additional_info = ""
    if additional_courses:
        additional_info = f"""
【同時申請中の他コース】
{', '.join(additional_courses)}
上記コースとの規定の矛盾・整合性も確認すること。
"""

    humax_info = ""
    if humax_knowledge:
        import json as _json
        humax_info = f"""
【ヒューマックス実務知識（支給要領に加えて必ずチェックすること）】
以下は社会保険労務士法人ヒューマックスが実務で蓄積した知識です。
支給要領ベースのチェックに加えて、以下の観点も必ずチェックしてください。

{_json.dumps(humax_knowledge, ensure_ascii=False, indent=2)}

上記の各項目について、就業規則に該当するリスクがあれば必ずアラートを出すこと。
"""

    law_info = ""
    if law_knowledge:
        import json as _json3
        law_info = f"""
【労働法令チェック知識（就業規則が法令に違反していないか確認すること）】
以下は労働基準法等の法令から抽出したチェックルールです。
支給要領・ヒューマックス知識に加えて、以下の法令違反がないかも必ずチェックしてください。

{_json3.dumps(law_knowledge, ensure_ascii=False, indent=2)}

上記の prohibited_patterns・common_violations に該当する記載があればCRITICALまたはWARNINGで指摘すること。
numeric_rules の数値要件を満たしていない場合もCRITICALで指摘すること。
"""

    ca_shoyo_text = ""
    if ca_shoyo_info:
        import json as _json2
        status    = ca_shoyo_info.get("status", "")
        shoyo     = ca_shoyo_info.get("shoyo", False)
        taishoku  = ca_shoyo_info.get("taishokukin", False)
        intro_date = ca_shoyo_info.get("introduction_date", "")
        targets   = []
        if shoyo:    targets.append("賞与")
        if taishoku: targets.append("退職金")

        ca_shoyo_text = f"""
【賞与・退職金制度導入コース（CA_shoyo）との同時申請情報】
- 導入予定制度：{"・".join(targets) if targets else "未選択"}
- 制度の状態：{status}{"（導入日：" + intro_date + "）" if intro_date else ""}

【CA_shoyo同時申請に基づく追加チェック指示】
"""
        if status == "制度導入前":
            if shoyo:
                ca_shoyo_text += """- 【CRITICAL必須】非正規雇用労働者（有期契約・パートタイム等）に賞与規定が適用されていないことを確認。適用されている場合は不支給要件に該当するためCRITICALで指摘すること。
"""
            if taishoku:
                ca_shoyo_text += """- 【CRITICAL必須】非正規雇用労働者（有期契約・パートタイム等）に退職金規定が適用されていないことを確認。適用されている場合は不支給要件に該当するためCRITICALで指摘すること。
- 【CRITICAL必須】退職金の支給額が「勤務月数×3,000円以上」の計算式になっているか確認。これはQ&Aに基づき正社員化コース単独申請でも判断基準として準用される。
"""
        elif status == "制度導入済み":
            ca_shoyo_text += f"""- 就業規則に導入日（{intro_date}）が正しく反映されているか確認。
- 導入後の規定内容が支給要領の要件を満たしているか確認。
"""
        if taishoku:
            ca_shoyo_text += """- 退職金の支給額が「勤務月数×3,000円以上」の計算式になっているか確認（制度導入前・後いずれの場合も）。
"""

    person_info_text = ""
    if target_person_info:
        person_info_text = f"""
【対象者情報（正社員化コース）】
※ この情報を元に、年齢・在籍期間・転換タイミングに関する要件を厳密にチェックすること。

- 生年月日：{target_person_info.get("生年月日", "不明")}（現在{target_person_info.get("年齢", "不明")}）
- 入社日：{target_person_info.get("入社日", "不明")}（在籍期間：{target_person_info.get("在籍期間", "不明")}）
- 転換予定日：{target_person_info.get("転換予定日", "不明")}（転換まで{target_person_info.get("転換まで", "不明")}）

【対象者情報に基づく必須チェック項目】
1. 転換予定日時点で6ヶ月以上の雇用期間を満たしているか
2. 雇用保険の加入要件（週所定労働時間20時間以上）を就業規則で確認できるか
3. 転換予定日から6ヶ月後の賃金支払いが確認できる規定か
4. 有期→無期→正規の転換ステップに問題はないか
5. 転換予定日時点の年齢と定年規定に矛盾がないか（定年まで6ヶ月以上あるか）
"""

    return f"""あなたは雇用関係助成金の申請業務に特化した社会保険労務士法人のベテラン監査官AIです。
就業規則を助成金の支給要領・Q&A・パンフレットに照らして厳格にチェックし、
不支給リスクをゼロにするための実務的な監査レポートを作成します。

【チェック対象の助成金】
{course_name}（{year}年度 {eff_date}施行版）

【チェックフェーズ】
{PHASE_LABELS.get(phase, phase)}

【判定の根拠知識】
以下のJSONが、このチェックで使用する支給要領・Q&A・パンフレットから抽出した判定ルールです。
このJSON以外の情報（インターネット上の情報等）は一切使用しないこと。

{rule_json}

【監査の姿勢】
- 「たぶん大丈夫」という曖昧な判断は絶対にしない
- 1円・1日・1文字でもズレていたら必ずアラートを出す
- グレーゾーンは「グレーゾーン」として明示し、人間確認を促す
- 不支給になった実例・審査官の目線で生々しく指摘する
- 修正案は条文レベルで具体的に提示する（コピペで使えるレベル）

【出力の制約】
- 問題がない項目（OK）は出力しない。問題がある項目のみ出力すること
- 推測や補完は行わない。記載がない場合は「記載なし」として指摘する
- アラートレベルは必ず以下のいずれかを使用すること：
  CRITICAL（不支給確定リスク）/ WARNING（不支給リスクあり）/ CAUTION（グレーゾーン）/ HUMAN_CHECK（人間確認必要）/ OK（問題なし）

【誤検知しやすいパターン（これらは問題なし・指摘不要）】
- 「ただし、有期契約社員、パート社員については別段の定めをしたときはその定めによる」という表現は標準的な就業規則の記載であり問題なし
- 「別段の定め」「別に定める」「別途定める」「別段の定めをしたときはその定めによる」という留保表現は問題なし。別規程がなくても問題なし。別規程が存在する場合のみ「再アップロードして確認を」とHUMAN_CHECKで案内する
- 正社員用規程にパート・有期の適用除外規定があること自体は問題なし（むしろ適切な記載）
- 「有期契約社員、パート社員については別段の定めをしたときはその定めによる」はCRITICALではなくOKまたはHUMAN_CHECK（別規程がある場合のみ確認）
- 正社員が日給制であること自体は問題なし。日給制のみで統一されている場合はWARNINGやCRITICALを出さないこと
- 正社員の給与形態は「月給制のみ」「日給制のみ」はOK。「月給制または日給制（混在）」は区分基準の記載がなければWARNING。「時給制」が正社員に適用されている場合のみCRITICAL
- 昇給規定に「毎年●月に行う」「原則として毎年●月に行う」と昇給月が明記されていれば定期昇給の要件を満たしているためOK。その後に「業績を勘案して各人ごとに決定する」「臨時昇給・臨時降給の規定がある」という記載があってもOK。昇給月の記載がある条文全体をNGと判断しないこと
- 正社員転換制度の条文について、転換要件として「勤続6ヶ月以上」「本人が希望する場合」「所属長の推薦」「面接試験合格」「正社員と同等の勤務時間・日数」などが記載されていればOK。転換時期が「随時」でもOK。客観性が不十分などの理由でWARNINGを出さないこと
- 法定通りの記載（労基法・育介法等の条文をそのまま引用したもの）は問題なし

【文章量の制約】★厳守★
- 各フィールドは簡潔に。ダラダラ書かない。
- 「問題の内容」：3行以内・150文字以内。何が問題かを一言で。
- 「該当箇所」：条文番号のリストのみ。説明不要。
- 「修正案」：「〜に修正する」「〜を追加する」など1〜2文で完結させること。例文は1つだけ。
- 「審査官の目線」：2行以内・100文字以内。
- 「根拠」：支給要領の条番号・項目名のみ。説明不要。
- 長い例文・丁寧な説明・背景の説明は不要。実務者が瞬時に判断できる情報だけを書く。

{phase_instruction}

{additional_info}
{humax_info}
{law_info}
{ca_shoyo_text}
{person_info_text}

【出力形式】
必ず以下の構造でMarkdownレポートを出力してください。

---

# 就業規則チェックレポート

## サマリー
- 🔴 CRITICAL（即時修正必須）: X件
- 🟡 WARNING（要修正）: X件
- 🟠 CAUTION（要確認）: X件
- 👤 HUMAN_CHECK（人間確認）: X件

---

## 詳細チェック結果

### [チェック項目名]

**判定：[アラートレベル] [アイコン]**

**該当箇所：**
（就業規則の条文番号・規程名・具体的な文言を引用）

**問題の内容：**
（3行以内・150文字以内で端的に。背景説明不要）

**審査官の目線：**
（2行以内・100文字以内で端的に）

**修正案：**
（1〜2文で完結。例文は1つまで）

**根拠：**
（条番号・項目名のみ）

---

（以降、全チェック項目を同じ形式で出力）

---

## 人間確認が必要な項目一覧
（HUMAN_CHECKの項目をまとめて再掲）

---

【就業規則本文】
{rules_text}
"""

# ============================================================
# チェック実行
# ============================================================

def run_audit(
    rules_text: str,
    rule_knowledge: dict,
    course_id: str,
    phase: str,
    additional_courses: list = None,
    target_person_info: dict = None,
    humax_knowledge: dict = None,
    ca_shoyo_info: dict = None,
    law_knowledge: dict = None,
) -> str:
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")
        if os.path.exists(env_path):
            with open(env_path, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if line.startswith("GOOGLE_API_KEY="):
                        api_key = line.split("=", 1)[1].strip()
                        break
    if not api_key:
        raise ValueError("GOOGLE_API_KEY が設定されていません。.env ファイルを確認してください。")

    client = genai.Client(api_key=api_key)

    prompt = build_prompt(rule_knowledge, course_id, phase, rules_text, additional_courses, target_person_info, humax_knowledge, ca_shoyo_info, law_knowledge)

    response = client.models.generate_content(
        model=MODEL,
        contents=prompt,
        config=types.GenerateContentConfig(
            max_output_tokens=MAX_TOKENS,
            temperature=0.1,
        ),
    )
    raw_report = response.text

    return raw_report

# ============================================================
# 就業規則サマリー抽出（正社員化コース専用）
# ============================================================

SEISHAIN_SUMMARY_ITEMS = [
    ("所定労働時間",         "1日・1週間の所定労働時間（例：1日8時間、週40時間）"),
    ("始業・終業・休憩時間", "始業時刻・終業時刻・休憩時間の具体的な時刻"),
    ("休日",                 "法定休日・所定休日（曜日・日数）"),
    ("変形労働時間制",       "適用の有無・種類（1ヶ月変形/1年変形/フレックス等）"),
    ("定年",                 "定年年齢・再雇用制度の有無"),
    ("給与形態",             "月給制/日給制/時給制の区分・基本給の決め方"),
    ("規定されている手当",   "手当名・支給条件・金額または計算方法の一覧"),
    ("昇給",                 "昇給の時期・基準・方法"),
    ("賞与",                 "支給時期・算定基準・支給月数の目安"),
    ("退職金",               "支給条件・算定方法・支給時期"),
]

def extract_seishain_summary(rules_text: str, api_key: str) -> list:
    """就業規則から正社員に関する主要事項を抽出してサマリーを生成"""
    client = genai.Client(api_key=api_key)

    items_desc = "\n".join(
        f"- {item}：{desc}" for item, desc in SEISHAIN_SUMMARY_ITEMS
    )

    prompt = f"""あなたは就業規則の専門家です。
以下の就業規則から、正社員（本則）に関する以下の項目を抽出してください。

【抽出項目】
{items_desc}

【抽出ルール】
- 就業規則に記載がある場合：条文番号と具体的な内容を簡潔に記載（50文字以内）
- 就業規則に記載がない場合：「記載なし」と記載
- 曖昧・不明確な場合：内容を記載した上で末尾に「※要確認」を付ける
- 複数の規程にまたがる場合は最も詳細な記載を使用

【出力形式】
必ず以下のJSON形式のみで出力すること。説明文は不要。

{{
  "所定労働時間": {{"条文": "第〇条", "内容": "〇〇", "要確認": false}},
  "始業・終業・休憩時間": {{"条文": "第〇条", "内容": "〇〇", "要確認": false}},
  "休日": {{"条文": "第〇条", "内容": "〇〇", "要確認": false}},
  "変形労働時間制": {{"条文": "第〇条", "内容": "〇〇", "要確認": false}},
  "定年": {{"条文": "第〇条", "内容": "〇〇", "要確認": false}},
  "給与形態": {{"条文": "第〇条", "内容": "〇〇", "要確認": false}},
  "規定されている手当": {{"条文": "第〇条〜", "内容": "〇〇", "要確認": false}},
  "昇給": {{"条文": "第〇条", "内容": "〇〇", "要確認": false}},
  "賞与": {{"条文": "第〇条", "内容": "〇〇", "要確認": false}},
  "退職金": {{"条文": "第〇条", "内容": "〇〇", "要確認": false}}
}}

【就業規則本文】
{rules_text[:30000]}
"""

    response = client.models.generate_content(
        model=MODEL,
        contents=prompt,
        config=types.GenerateContentConfig(
            max_output_tokens=4096,
            temperature=0.0,
        ),
    )

    import json, re
    raw = response.text.strip()
    # JSONブロックを抽出
    m = re.search(r'\{[\s\S]+\}', raw)
    if m:
        try:
            return json.loads(m.group())
        except json.JSONDecodeError:
            pass
    return {}


# ============================================================
# レポート生成
# ============================================================

def generate_report(
    audit_result: str,
    company_name: str,
    course_id: str,
    phase: str,
    rule_knowledge: dict,
    target_files: list,
) -> str:
    course_name = rule_knowledge.get("meta", {}).get("course_name", course_id)
    year = rule_knowledge.get("meta", {}).get("year", "")
    eff_date = rule_knowledge.get("meta", {}).get("effective_date", "")
    now = datetime.now().strftime("%Y年%m月%d日 %H:%M")

    header = f"""# 就業規則チェックレポート

| 項目 | 内容 |
|------|------|
| 会社名 | {company_name} |
| チェック日時 | {now} |
| 対象助成金 | {course_name} |
| 適用支給要領 | {year}年度 {eff_date}施行版 |
| チェックフェーズ | {PHASE_LABELS.get(phase, phase)} |
| チェック対象ファイル | {', '.join(target_files)} |

---

"""
    result_body = audit_result
    if result_body.startswith("# 就業規則チェックレポート"):
        lines = result_body.split("\n")
        result_body = "\n".join(lines[1:]).lstrip("\n")

    return header + result_body
